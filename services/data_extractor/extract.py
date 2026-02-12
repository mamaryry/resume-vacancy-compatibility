"""
Resume text extraction from PDF and DOCX files.

This module provides functions to extract text content from various resume file formats,
with robust error handling for malformed files.
"""
import logging
from pathlib import Path
from typing import Dict, Optional, Union

import pdfplumber
from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_pdf(
    file_path: Union[str, Path], use_fallback: bool = True
) -> Dict[str, Optional[str]]:
    """
    Extract text from a PDF file using PyPDF2 and pdfplumber as fallback.

    Args:
        file_path: Path to the PDF file
        use_fallback: If True, try pdfplumber if PyPDF2 fails or returns minimal text

    Returns:
        Dictionary containing:
            - text: Extracted text content (None if extraction fails)
            - method: Which library succeeded ('pypdf2', 'pdfplumber', or None)
            - pages: Number of pages detected
            - error: Error message if extraction failed

    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If the file is not a valid PDF

    Examples:
        >>> result = extract_text_from_pdf("resume.pdf")
        >>> print(result["text"])
        'John Doe\\nSoftware Engineer...'
        >>> print(result["method"])
        'pypdf2'
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if not file_path.suffix.lower() == ".pdf":
        raise ValueError(f"File is not a PDF: {file_path}")

    # Try PyPDF2 first (faster)
    try:
        result = _extract_with_pypdf2(file_path)
        # Check if we got meaningful content
        text_length = len(result["text"].strip()) if result["text"] else 0

        if text_length > 50 or not use_fallback:
            logger.info(f"Extracted {text_length} chars from {file_path.name} using PyPDF2")
            return result
        else:
            logger.warning(
                f"PyPDF2 extracted minimal text ({text_length} chars), trying pdfplumber"
            )
    except Exception as e:
        logger.warning(f"PyPDF2 extraction failed: {e}")
        if not use_fallback:
            return {
                "text": None,
                "method": None,
                "pages": 0,
                "error": f"PyPDF2 failed: {str(e)}",
            }

    # Fallback to pdfplumber (better for complex layouts)
    if use_fallback:
        try:
            result = _extract_with_pdfplumber(file_path)
            text_length = len(result["text"].strip()) if result["text"] else 0
            logger.info(
                f"Extracted {text_length} chars from {file_path.name} using pdfplumber"
            )
            return result
        except Exception as e:
            logger.error(f"pdfplumber extraction also failed: {e}")
            return {
                "text": None,
                "method": None,
                "pages": 0,
                "error": f"All extraction methods failed: {str(e)}",
            }

    return {
        "text": None,
        "method": None,
        "pages": 0,
        "error": "No extraction method succeeded",
    }


def _extract_with_pypdf2(file_path: Path) -> Dict[str, Optional[str]]:
    """
    Extract text using PyPDF2 library.

    Args:
        file_path: Path to the PDF file

    Returns:
        Dictionary with extracted text and metadata
    """
    try:
        reader = PdfReader(str(file_path))
        num_pages = len(reader.pages)

        text_parts = []
        for page_num, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"Failed to extract page {page_num}: {e}")
                continue

        text = "\n\n".join(text_parts) if text_parts else ""

        return {
            "text": text if text.strip() else None,
            "method": "pypdf2",
            "pages": num_pages,
            "error": None,
        }

    except Exception as e:
        raise RuntimeError(f"PyPDF2 extraction error: {e}") from e


def _extract_with_pdfplumber(file_path: Path) -> Dict[str, Optional[str]]:
    """
    Extract text using pdfplumber library.

    Pdfplumber is more robust for complex PDF layouts and handles
    some edge cases better than PyPDF2.

    Args:
        file_path: Path to the PDF file

    Returns:
        Dictionary with extracted text and metadata
    """
    try:
        with pdfplumber.open(file_path) as pdf:
            num_pages = len(pdf.pages)
            text_parts = []

            for page_num, page in enumerate(pdf.pages, start=1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"pdfplumber failed to extract page {page_num}: {e}")
                    continue

            text = "\n\n".join(text_parts) if text_parts else ""

            return {
                "text": text if text.strip() else None,
                "method": "pdfplumber",
                "pages": num_pages,
                "error": None,
            }

    except Exception as e:
        raise RuntimeError(f"pdfplumber extraction error: {e}") from e


def validate_pdf_file(file_path: Union[str, Path]) -> Dict[str, Union[bool, str]]:
    """
    Validate a PDF file before extraction.

    Checks:
    - File exists
    - Has .pdf extension
    - Is not empty
    - Can be opened by PDF libraries

    Args:
        file_path: Path to the PDF file

    Returns:
        Dictionary with validation results:
            - valid: Boolean indicating if file is valid
            - reason: String explaining why validation failed (if applicable)

    Examples:
        >>> validation = validate_pdf_file("resume.pdf")
        >>> if validation["valid"]:
        ...     result = extract_text_from_pdf("resume.pdf")
    """
    file_path = Path(file_path)

    # Check existence
    if not file_path.exists():
        return {"valid": False, "reason": "File not found"}

    # Check extension
    if file_path.suffix.lower() != ".pdf":
        return {"valid": False, "reason": "Not a PDF file"}

    # Check file size
    if file_path.stat().st_size == 0:
        return {"valid": False, "reason": "File is empty"}

    # Try to open with PyPDF2
    try:
        reader = PdfReader(str(file_path))
        if len(reader.pages) == 0:
            return {"valid": False, "reason": "PDF has no pages"}
    except Exception as e:
        return {"valid": False, "reason": f"Cannot open PDF: {str(e)}"}

    return {"valid": True, "reason": None}


def extract_text_from_docx(
    file_path: Union[str, Path]
) -> Dict[str, Optional[str]]:
    """
    Extract text from a DOCX file using python-docx.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Dictionary containing:
            - text: Extracted text content (None if extraction fails)
            - method: Always 'python-docx' if successful
            - paragraphs: Number of paragraphs extracted
            - error: Error message if extraction failed

    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If the file is not a valid DOCX

    Examples:
        >>> result = extract_text_from_docx("resume.docx")
        >>> print(result["text"])
        'John Doe\\nSoftware Engineer...'
        >>> print(result["method"])
        'python-docx'
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if not file_path.suffix.lower() in [".docx", ".doc"]:
        raise ValueError(f"File is not a DOCX: {file_path}")

    try:
        result = _extract_with_python_docx(file_path)
        text_length = len(result["text"].strip()) if result["text"] else 0
        logger.info(
            f"Extracted {text_length} chars from {file_path.name} using python-docx"
        )
        return result
    except Exception as e:
        logger.error(f"python-docx extraction failed: {e}")
        return {
            "text": None,
            "method": None,
            "paragraphs": 0,
            "error": f"DOCX extraction failed: {str(e)}",
        }


def _extract_with_python_docx(file_path: Path) -> Dict[str, Optional[str]]:
    """
    Extract text using python-docx library.

    This function extracts all paragraphs from a DOCX document and joins them
    with newlines. It preserves the document structure by reading paragraphs
    in order.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Dictionary with extracted text and metadata

    Raises:
        RuntimeError: If extraction fails
    """
    try:
        doc = Document(str(file_path))

        # Extract all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():  # Only include non-empty paragraphs
                paragraphs.append(para.text.strip())

        # Join paragraphs with double newlines
        text = "\n\n".join(paragraphs) if paragraphs else ""

        # Also extract table contents (many resumes use tables)
        tables_text = _extract_tables_from_docx(doc)
        if tables_text:
            text = text + "\n\n" + tables_text if text else tables_text

        return {
            "text": text if text.strip() else None,
            "method": "python-docx",
            "paragraphs": len(doc.paragraphs),
            "error": None,
        }

    except Exception as e:
        raise RuntimeError(f"python-docx extraction error: {e}") from e


def _extract_tables_from_docx(doc: Document) -> str:
    """
    Extract text from all tables in a DOCX document.

    Many resume templates use tables for layout, so we need to extract
    table contents to get the complete text.

    Args:
        doc: python-docx Document object

    Returns:
        String containing all table contents, formatted with newlines
    """
    table_parts = []

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                table_parts.append(" | ".join(row_text))

    return "\n".join(table_parts) if table_parts else ""


def validate_docx_file(file_path: Union[str, Path]) -> Dict[str, Union[bool, str]]:
    """
    Validate a DOCX file before extraction.

    Checks:
    - File exists
    - Has .docx or .doc extension
    - Is not empty
    - Can be opened by python-docx

    Args:
        file_path: Path to the DOCX file

    Returns:
        Dictionary with validation results:
            - valid: Boolean indicating if file is valid
            - reason: String explaining why validation failed (if applicable)

    Examples:
        >>> validation = validate_docx_file("resume.docx")
        >>> if validation["valid"]:
        ...     result = extract_text_from_docx("resume.docx")
    """
    file_path = Path(file_path)

    # Check existence
    if not file_path.exists():
        return {"valid": False, "reason": "File not found"}

    # Check extension
    if file_path.suffix.lower() not in [".docx", ".doc"]:
        return {"valid": False, "reason": "Not a DOCX file"}

    # Check file size
    if file_path.stat().st_size == 0:
        return {"valid": False, "reason": "File is empty"}

    # Try to open with python-docx
    try:
        doc = Document(str(file_path))
        # Basic validation: check if document has any content
        if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
            return {"valid": False, "reason": "DOCX has no content"}
    except Exception as e:
        return {"valid": False, "reason": f"Cannot open DOCX: {str(e)}"}

    return {"valid": True, "reason": None}
