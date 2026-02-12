"""
Create sample test files for PDF and DOCX extraction testing.
"""
import sys
from pathlib import Path

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from docx import Document

    # Create sample PDF
    pdf_path = Path(__file__).parent / "sample.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)

    # Add content to PDF (resume-like content)
    c.drawString(100, 750, "JOHN DOE")
    c.drawString(100, 730, "Software Engineer")
    c.drawString(100, 700, "")

    c.drawString(100, 680, "SUMMARY")
    c.drawString(100, 660, "Experienced software engineer with 5+ years in full-stack development.")

    c.drawString(100, 630, "EXPERIENCE")
    c.drawString(100, 610, "Senior Developer | Tech Corporation | 2020 - Present")
    c.drawString(100, 590, "- Led development of microservices architecture")
    c.drawString(100, 570, "- Implemented CI/CD pipelines")
    c.drawString(100, 550, "- Mentored junior developers")

    c.drawString(100, 520, "Junior Developer | Startup Inc | 2018 - 2020")
    c.drawString(100, 500, "- Developed React-based web applications")
    c.drawString(100, 480, "- Built RESTful APIs with Python/Django")

    c.drawString(100, 450, "SKILLS")
    c.drawString(100, 430, "Programming: Python, JavaScript, TypeScript, Java")
    c.drawString(100, 410, "Frameworks: React, Spring Boot, Django")
    c.drawString(100, 390, "Databases: PostgreSQL, MongoDB, Redis")
    c.drawString(100, 370, "Tools: Docker, Kubernetes, Git, AWS")

    c.drawString(100, 340, "EDUCATION")
    c.drawString(100, 320, "Bachelor of Science in Computer Science")
    c.drawString(100, 300, "State University | 2014 - 2018")

    c.save()

    print(f"Created sample PDF: {pdf_path}")

    # Create sample DOCX
    docx_path = Path(__file__).parent / "sample.docx"
    doc = Document()

    doc.add_heading("JANE SMITH", 0)
    doc.add_paragraph("Data Scientist")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "Data scientist with expertise in machine learning, statistical analysis, "
        "and data visualization. Passionate about turning complex data into actionable insights."
    )

    doc.add_heading("Experience", level=1)

    p1 = doc.add_paragraph()
    p1.add_run("Senior Data Scientist ").bold = True
    p1.add_run("| Data Corp | 2021 - Present")
    doc.add_paragraph("- Developed predictive models using Python and scikit-learn", style='List Bullet')
    doc.add_paragraph("- Implemented NLP pipelines for text analysis", style='List Bullet')
    doc.add_paragraph("- Led team of 3 data analysts", style='List Bullet')

    p2 = doc.add_paragraph()
    p2.add_run("Data Analyst ").bold = True
    p2.add_run("| Analytics Inc | 2019 - 2021")
    doc.add_paragraph("- Created dashboards using Tableau and Power BI", style='List Bullet')
    doc.add_paragraph("- Automated data processing workflows", style='List Bullet')

    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Programming: Python, R, SQL, MATLAB", style='List Bullet')
    doc.add_paragraph("ML/AI: TensorFlow, PyTorch, scikit-learn, Keras", style='List Bullet')
    doc.add_paragraph("Visualization: Tableau, Power BI, Matplotlib, Plotly", style='List Bullet')
    doc.add_paragraph("Databases: PostgreSQL, MySQL, MongoDB", style='List Bullet')

    doc.add_heading("Education", level=1)
    doc.add_paragraph("Master of Science in Data Science")
    doc.add_paragraph("Tech University | 2017 - 2019")

    doc.add_paragraph("Bachelor of Science in Statistics")
    doc.add_paragraph("State University | 2013 - 2017")

    doc.save(str(docx_path))

    print(f"Created sample DOCX: {docx_path}")

    # Create a DOCX with tables (common resume format)
    docx_tables_path = Path(__file__).parent / "sample_with_tables.docx"
    doc = Document()

    doc.add_heading("RESUME: ROBERT JOHNSON", 0)

    # Add a table with contact info
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    table.rows[0].cells[0].text = "Email:"
    table.rows[0].cells[1].text = "robert.johnson@email.com"
    table.rows[1].cells[0].text = "Phone:"
    table.rows[1].cells[1].text = "+1 (555) 123-4567"
    table.rows[2].cells[0].text = "Location:"
    table.rows[2].cells[1].text = "San Francisco, CA"
    table.rows[3].cells[0].text = "LinkedIn:"
    table.rows[3].cells[1].text = "linkedin.com/in/robertjohnson"

    doc.add_paragraph()  # spacer

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "DevOps Engineer with 6+ years of experience in cloud infrastructure, "
        "automation, and CI/CD pipelines. Proven track record of reducing deployment "
        "time by 80% through automation."
    )

    doc.add_heading("Work Experience", level=1)

    # Experience table
    exp_table = doc.add_table(rows=4, cols=2)
    exp_table.style = 'Table Grid'

    exp_table.rows[0].cells[0].text = "Position:"
    exp_table.rows[0].cells[1].text = "DevOps Engineer"
    exp_table.rows[1].cells[0].text = "Company:"
    exp_table.rows[1].cells[1].text = "CloudTech Solutions"
    exp_table.rows[2].cells[0].text = "Duration:"
    exp_table.rows[2].cells[1].text = "2020 - Present"
    exp_table.rows[3].cells[0].text = "Key Achievements:"
    exp_table.rows[3].cells[1].text = (
        "• Designed and maintained AWS infrastructure\n"
        "• Implemented GitLab CI/CD pipelines\n"
        "• Reduced deployment time from 2 hours to 20 minutes\n"
        "• Managed Kubernetes clusters"
    )

    doc.add_paragraph()

    doc.add_heading("Technical Skills", level=1)

    # Skills table
    skills_table = doc.add_table(rows=5, cols=2)
    skills_table.style = 'Table Grid'

    skills_table.rows[0].cells[0].text = "Cloud Platforms:"
    skills_table.rows[0].cells[1].text = "AWS, GCP, Azure"
    skills_table.rows[1].cells[0].text = "Containers:"
    skills_table.rows[1].cells[1].text = "Docker, Kubernetes, ECS"
    skills_table.rows[2].cells[0].text = "CI/CD:"
    skills_table.rows[2].cells[1].text = "Jenkins, GitLab CI, GitHub Actions"
    skills_table.rows[3].cells[0].text = "Infrastructure as Code:"
    skills_table.rows[3].cells[1].text = "Terraform, CloudFormation, Ansible"
    skills_table.rows[4].cells[0].text = "Monitoring:"
    skills_table.rows[4].cells[1].text = "Prometheus, Grafana, ELK Stack"

    doc.save(str(docx_tables_path))

    print(f"Created sample DOCX with tables: {docx_tables_path}")

    print("\n✓ All sample files created successfully!")

except ImportError as e:
    print(f"Error: {e}")
    print("\nPlease install required packages:")
    print("  pip install reportlab python-docx")
    sys.exit(1)
